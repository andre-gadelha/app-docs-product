from datetime import datetime, timedelta
from copy import deepcopy
import os
import re
import shutil
import uuid
from pathlib import Path
import xml.etree.ElementTree as ET
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from flask import current_app
from werkzeug.utils import secure_filename

class DocxService:
    XLSX_NS = {'a': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    W_TAG = '{%s}' % W_NS

    def _resolve_template_path(self, config_key):
        template_path = current_app.config.get(config_key)
        if template_path and os.path.exists(template_path):
            return template_path
        raise FileNotFoundError(
            f"Template não encontrado no caminho configurado: {template_path}"
        )

    def _resolve_data_source_file(self):
        base_dir = Path(current_app.config['FONTE_DADOS_DIR'])
        if not base_dir.exists():
            raise FileNotFoundError(f"Pasta de fonte de dados não encontrada: {base_dir}")
        xlsx_files = sorted(base_dir.glob('*.xlsx'))
        if not xlsx_files:
            raise FileNotFoundError(f"Nenhuma planilha .xlsx encontrada em: {base_dir}")
        return xlsx_files[0]

    def _insert_after_paragraph(self, paragraph, text, align=None):
        new_p_elm = OxmlElement('w:p')
        paragraph._p.addnext(new_p_elm)
        new_p = Paragraph(new_p_elm, paragraph._parent)
        if text:
            new_p.add_run(text)
        if align is not None:
            new_p.alignment = align
        return new_p

    def _ensure_runtime_folders(self):
        folders = [
            current_app.config['UPLOAD_FOLDER'],
            current_app.config['UPLOAD_PROPOSTAS_FOLDER'],
            current_app.config['UPLOAD_HUS_FOLDER'],
            current_app.config['UPLOAD_RELATORIOS_FOLDER'],
        ]
        for folder in folders:
            Path(folder).mkdir(parents=True, exist_ok=True)

    def _new_execution_paths(self):
        execution_id = datetime.now().strftime('%Y%m%d%H%M%S') + '_' + uuid.uuid4().hex[:8]
        proposta_dir = Path(current_app.config['UPLOAD_PROPOSTAS_FOLDER']) / execution_id
        hu_dir = Path(current_app.config['UPLOAD_HUS_FOLDER']) / execution_id
        relatorio_dir = Path(current_app.config['UPLOAD_RELATORIOS_FOLDER']) / execution_id
        proposta_dir.mkdir(parents=True, exist_ok=True)
        hu_dir.mkdir(parents=True, exist_ok=True)
        relatorio_dir.mkdir(parents=True, exist_ok=True)
        return proposta_dir, hu_dir, relatorio_dir

    def _apply_inter_12_to_paragraph(self, paragraph):
        """Formatação padrão do relatório (lista § 4.1 e alinhamentos ao template)."""
        if paragraph._p.getparent() is None:
            return
        for run in paragraph.runs:
            run.font.name = 'Inter'
            run.font.size = Pt(12)

    def _replace_text_preserve_format(self, paragraph, key, value):
        if key not in paragraph.text:
            return
        if paragraph.runs:
            full_text = paragraph.text.replace(key, value)
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ''
            paragraph.runs[0].font.name = 'Inter'
        else:
            paragraph.text = paragraph.text.replace(key, value)

    def _replace_placeholders_everywhere(self, doc, replacements):
        for p in doc.paragraphs:
            for key, value in replacements.items():
                self._replace_text_preserve_format(p, key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in replacements.items():
                            self._replace_text_preserve_format(p, key, value)

        for section in doc.sections:
            for header in [section.header, section.footer]:
                for p in header.paragraphs:
                    for key, value in replacements.items():
                        self._replace_text_preserve_format(p, key, value)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for key, value in replacements.items():
                                    self._replace_text_preserve_format(p, key, value)

    def _format_brl_value(self, value):
        """Formata valor numérico no padrão brasileiro sem prefixo de moeda."""
        formatted = f"{value:,.2f}"
        return formatted.replace(",", "X").replace(".", ",").replace("X", ".")

    def _extract_intro_from_proposta_docx(self, proposta_path):
        proposta_doc = Document(str(proposta_path))
        paragraphs = [p.text.strip() for p in proposta_doc.paragraphs if p.text and p.text.strip()]
        if not paragraphs:
            return ''
        for idx, text in enumerate(paragraphs):
            lowered = text.casefold()
            if lowered == 'introducao' or lowered == 'introdução':
                if idx + 1 < len(paragraphs):
                    return paragraphs[idx + 1]
        return paragraphs[0]

    def _extract_proposta_identity(self, proposta_path, original_filename):
        filename = Path(original_filename or proposta_path.name).name
        filename_match = re.search(
            r"(?i)proposta\s*os\s*(\d{4})\s*-\s*([0-9]+)\s*-\s*(.+?)\.docx$",
            filename
        )
        if filename_match:
            return {
                'ano_os': filename_match.group(1),
                'os': filename_match.group(2).zfill(4),
                'assunto': filename_match.group(3).strip(),
            }

        proposta_doc = Document(str(proposta_path))
        all_text = '\n'.join([p.text for p in proposta_doc.paragraphs if p.text]).strip()
        text_match = re.search(r"(?i)(\d{4})\s*-\s*([0-9]{1,4})", all_text)
        if text_match:
            return {
                'ano_os': text_match.group(1),
                'os': text_match.group(2).zfill(4),
                'assunto': '',
            }
        raise ValueError(
            "Não foi possível identificar Ano OS e Número OS a partir da proposta enviada. "
            "Use um arquivo no padrão: Proposta OS 2026-0032 - Assunto.docx"
        )

    def _normalize_header(self, header):
        key = (header or '').strip().casefold()
        aliases = {
            'ano os': 'ano_os',
            'ano_os': 'ano_os',
            'os': 'os',
            'assunto': 'assunto',
            'início': 'inicio',
            'inicio': 'inicio',
            'fim': 'fim',
            'hsts': 'hst',
            'hst': 'hst',
            'valor': 'valor',
        }
        return aliases.get(key, key)

    def _read_xlsx_row_by_os(self, workbook_path, ano_os, os_number):
        with ZipFile(workbook_path) as archive:
            shared = []
            if 'xl/sharedStrings.xml' in archive.namelist():
                root = ET.fromstring(archive.read('xl/sharedStrings.xml'))
                for item in root.findall('a:si', self.XLSX_NS):
                    shared.append(''.join((t.text or '') for t in item.findall('.//a:t', self.XLSX_NS)))

            sheet = ET.fromstring(archive.read('xl/worksheets/sheet1.xml'))
            rows = []
            for row in sheet.findall('a:sheetData/a:row', self.XLSX_NS):
                values = {}
                for cell in row.findall('a:c', self.XLSX_NS):
                    ref = cell.attrib.get('r', 'A1')
                    col_letters = ''.join(ch for ch in ref if ch.isalpha())
                    index = 0
                    for ch in col_letters:
                        index = index * 26 + ord(ch.upper()) - 64
                    col_index = index - 1
                    cell_type = cell.attrib.get('t')
                    raw_node = cell.find('a:v', self.XLSX_NS)
                    raw = '' if raw_node is None or raw_node.text is None else raw_node.text
                    if cell_type == 'inlineStr':
                        value = ''.join((t.text or '') for t in cell.findall('.//a:t', self.XLSX_NS))
                    elif cell_type == 's' and raw:
                        value = shared[int(raw)]
                    else:
                        value = raw
                    values[col_index] = value
                rows.append(values)

        if not rows:
            raise ValueError("Planilha sem conteúdo.")

        max_col = max(rows[0].keys()) if rows[0] else -1
        headers = []
        for i in range(max_col + 1):
            headers.append(rows[0].get(i, '').strip())
        normalized = [self._normalize_header(h) for h in headers]

        for values in rows[1:]:
            row_data = {}
            for i, key in enumerate(normalized):
                row_data[key] = str(values.get(i, '')).strip()
            row_ano = row_data.get('ano_os', '')
            row_os = row_data.get('os', '').zfill(4)
            if row_ano == str(ano_os) and row_os == str(os_number).zfill(4):
                return row_data
        raise ValueError(f"OS {ano_os}-{str(os_number).zfill(4)} não encontrada na planilha.")

    def _excel_serial_to_date(self, value):
        if not value:
            return ''
        try:
            serial = float(str(value).replace(',', '.'))
            dt = datetime(1899, 12, 30) + timedelta(days=serial)
            return dt.strftime('%d/%m/%Y')
        except Exception:
            return str(value)

    def _paragraph_matches_token(self, text, tokens):
        normalized = (text or '').strip().replace('ç', 'c').replace('Ç', 'C').replace('ã', 'a').replace('Ã', 'A')
        return any(token in normalized for token in tokens)

    def _insert_page_break_before(self, paragraph):
        if paragraph._p.getparent() is None:
            return
        break_p_elm = OxmlElement('w:p')
        break_r = OxmlElement('w:r')
        break_br = OxmlElement('w:br')
        break_br.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'page')
        break_r.append(break_br)
        break_p_elm.append(break_r)
        paragraph._p.addprevious(break_p_elm)

    def _force_p_center_ooxml(self, p_elm):
        p_pr = p_elm.find(self.W_TAG + 'pPr')
        if p_pr is None:
            p_pr = OxmlElement('w:pPr')
            p_elm.insert(0, p_pr)
        jc = p_pr.find(self.W_TAG + 'jc')
        if jc is None:
            jc = OxmlElement('w:jc')
            p_pr.append(jc)
        jc.set(qn('w:val'), 'center')

    def _ct_p_plain_text_from_elm(self, p_elm):
        parts = []
        for t in p_elm.iter(self.W_TAG + 't'):
            if t.text:
                parts.append(t.text)
        return ''.join(parts).strip()

    def _paragraph_clear_top_spacing_ooxml(self, p_elm):
        p_pr = p_elm.find(self.W_TAG + 'pPr')
        if p_pr is None:
            return
        sp = p_pr.find(self.W_TAG + 'spacing')
        if sp is None:
            return
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:beforeLines'), '0')

    def _remove_preceding_empty_paragraphs(self, paragraph):
        par = paragraph._p
        parent = par.getparent()
        if parent is None:
            return
        while True:
            prev = par.getprevious()
            if prev is None or prev.tag != self.W_TAG + 'p':
                break
            if self._ct_p_plain_text_from_elm(prev):
                break
            parent.remove(prev)

    def _item9_anexos_heading_re(self):
        return re.compile(r'^\s*9\.\s*Anexo', re.IGNORECASE)

    def _find_item9_anexos_heading_paragraph(self, doc):
        """Parágrafo do título «9. Anexos» (ou «9. Anexo») no corpo ou em células de tabela."""
        heading_re = self._item9_anexos_heading_re()

        def _scan(paragraphs):
            for p in paragraphs:
                if heading_re.search(p.text or ''):
                    return p
            return None

        p = _scan(doc.paragraphs)
        if p:
            return p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    p = _scan(cell.paragraphs)
                    if p:
                        return p
        return None

    def _normalize_item9_anexos_heading_position(self, doc):
        """Coloca «9. Anexos» no topo útil da página (sem espaço extra nem parágrafos vazios acima)."""
        p = self._find_item9_anexos_heading_paragraph(doc)
        if not p:
            return
        try:
            p.paragraph_format.space_before = Pt(0)
        except Exception:
            pass
        self._paragraph_clear_top_spacing_ooxml(p._p)
        self._remove_preceding_empty_paragraphs(p)

    def _ensure_item9_anexos_starts_new_page(self, doc):
        """«9. Anexos» inicia na primeira linha da página seguinte ao § 8."""
        p = self._find_item9_anexos_heading_paragraph(doc)
        if p and p._p.getparent() is not None:
            self._insert_page_break_before(p)

    def _ensure_two_blank_lines_after_item9_heading(self, doc):
        """Duas linhas em branco entre o título do § 9 e a primeira linha da lista (ANEXO 1)."""
        p = self._find_item9_anexos_heading_paragraph(doc)
        if not p or p._p.getparent() is None:
            return
        cur = p
        for _ in range(2):
            cur = self._insert_after_paragraph(cur, '')

    def _content_hu_placeholder_pattern(self):
        # Template: {{Conteúdo da HU correspondente na lista de anexos}} ou ...lista de anexo}}
        return re.compile(
            r'\{\{\s*Conte[uú]do\s+da\s+HU\s+correspondente\s+na\s+lista\s+de\s+anexo[s]?\s*\}\}',
            re.IGNORECASE,
        )

    def _collect_paragraphs_with_pattern(self, doc, pattern):
        found = []
        for paragraph in doc.paragraphs:
            if pattern.search(paragraph.text or ''):
                found.append(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if pattern.search(paragraph.text or ''):
                            found.append(paragraph)
        for section in doc.sections:
            for header in (section.header, section.footer):
                for paragraph in header.paragraphs:
                    if pattern.search(paragraph.text or ''):
                        found.append(paragraph)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if pattern.search(paragraph.text or ''):
                                    found.append(paragraph)
        return found

    def _write_hu_item_9(self, doc, hu_items):
        """
        Item 9: lista de anexos alinhada à secção 4.1 — um placeholder por HU,
        cada um numa nova página, só com «ANEXO N - {rótulo do upload}» centrado.
        """
        pattern = self._content_hu_placeholder_pattern()
        placeholder_paragraphs = self._collect_paragraphs_with_pattern(doc, pattern)

        if not placeholder_paragraphs:
            return

        template_p_xml = deepcopy(placeholder_paragraphs[0]._p)
        anchor_tail = placeholder_paragraphs[-1]
        while len(placeholder_paragraphs) < len(hu_items):
            cloned = deepcopy(template_p_xml)
            anchor_tail._p.addnext(cloned)
            new_p = Paragraph(cloned, anchor_tail._parent)
            placeholder_paragraphs.append(new_p)
            anchor_tail = new_p

        for j in range(len(hu_items), len(placeholder_paragraphs)):
            p = placeholder_paragraphs[j]
            p.text = pattern.sub('', p.text or '').strip()

        for idx, hu_item in enumerate(hu_items):
            paragraph = placeholder_paragraphs[idx]
            if paragraph._p.getparent() is None:
                continue

            if idx > 0:
                self._insert_page_break_before(paragraph)

            paragraph.text = pattern.sub('', paragraph.text or '').strip()
            label = (hu_item.get('label') or '').strip() or f'HU {idx + 1}'
            line = f'ANEXO {idx + 1} - {label}'
            paragraph.text = line
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._force_p_center_ooxml(paragraph._p)

    def _write_hu_list_item_41(self, doc, hu_items):
        pattern = re.compile(r'\{\{\s*HU\s+(\d+|N)\s*\}\}', re.IGNORECASE)
        placeholders = [p for p in doc.paragraphs if pattern.search(p.text or '')]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if pattern.search(p.text or ''):
                            placeholders.append(p)

        if placeholders:
            template_p = next(
                (p for p in placeholders if p._p.pPr is not None and p._p.pPr.numPr is not None),
                placeholders[0]
            )
            template_xml = deepcopy(template_p._p)
            anchor_tail = placeholders[-1]
            while len(placeholders) < len(hu_items):
                cloned = deepcopy(template_xml)
                anchor_tail._p.addnext(cloned)
                p_new = Paragraph(cloned, anchor_tail._parent)
                placeholders.append(p_new)
                anchor_tail = p_new
            for idx, hu in enumerate(hu_items):
                p = placeholders[idx]
                label = hu['label']
                p_text = p.text or ''
                matched = pattern.search(p_text)
                if matched:
                    p_text = p_text.replace(matched.group(0), label)
                else:
                    p_text = label
                # Linhas clonadas repetem «Anexo 1» do modelo; alinha à ordem do formulário.
                p_text = re.sub(
                    r'(?i)Anexo\s*\d+',
                    f'Anexo {idx + 1}',
                    p_text,
                    count=1,
                )
                p.text = p_text
                self._apply_inter_12_to_paragraph(p)
            for idx in range(len(hu_items), len(placeholders)):
                placeholders[idx].text = ''
                self._apply_inter_12_to_paragraph(placeholders[idx])
            return

        for paragraph in doc.paragraphs:
            if '<Lista HUs>' in paragraph.text or '{{HU_LISTA}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('<Lista HUs>', '').replace('{{HU_LISTA}}', '')
                current = paragraph
                for idx, hu in enumerate(hu_items):
                    if idx == 0:
                        current.text = f"Anexo {idx + 1}: {hu['label']}"
                        current.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        self._apply_inter_12_to_paragraph(current)
                    else:
                        current = self._insert_after_paragraph(current, f"Anexo {idx + 1}: {hu['label']}")
                        current.style = paragraph.style
                        current.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        self._apply_inter_12_to_paragraph(current)
                return

    def _sanitize_upload_name(self, filename, fallback):
        clean = secure_filename(filename or '')
        return clean or fallback

    def _sanitize_output_filename(self, text):
        sanitized = re.sub(r'[<>:"/\\|?*]', '-', text).strip().rstrip('.')
        return sanitized or 'Relatório de Entrega'

    def generate_proposta_os(self, data):
        self._ensure_runtime_folders()
        template_path = self._resolve_template_path('TEMPLATE_DOCX')

        doc = Document(template_path)

        # Dados para substituição
        ano_os = data.get('ano_os', '')
        num_os = data.get('num_os', '')
        titulo_os = data.get('titulo_os', '')
        nome_os = data.get('nome_os', '')
        tipo_os = data.get('tipo_os', '')
        nome_autor = data.get('nome_autor', '')
        nome_solicitante = data.get('nome_solicitante', '')
        descricao_geral = data.get('descricao_geral', '')
        qtd_hst = int(float(data.get('qtd_hst', 0))) # Formatação sem casas decimais
        itens = data.get('itens', [])
        data_atual = datetime.now().strftime('%d/%m/%Y')
        valor_calculado = self._format_brl_value(qtd_hst * 200)

        replacements = {
            '<Autor>': nome_autor,
            '<Nome da OS>': nome_os,
            '<Nome da Os>': nome_os,
            '<Tipo da OS>': tipo_os,
            '<Solicitante>': nome_solicitante,
            '<Nome do solicitante>': nome_solicitante,
            '<Descrição Geral da OS>': descricao_geral,
            '<Quantidade de HST>': str(qtd_hst),
            '<Data atual do servidor>': data_atual,
            '<Cálculo em R$ do valor de HST x 200,00>': valor_calculado
        }

        self._replace_placeholders_everywhere(doc, replacements)

        # Tratamento especial para <Itens da OS>
        for p in doc.paragraphs:
            if '<Itens da OS>' in p.text:
                itens_text = ""
                for i, item in enumerate(itens, 1):
                    itens_text += f"{i}. {item}\n"
                # Para itens, substituímos o texto e forçamos alinhamento, mas mantemos o estilo do parágrafo
                p.text = p.text.replace('<Itens da OS>', itens_text.strip())
                for run in p.runs:
                    run.font.name = 'Inter'
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '<Itens da OS>' in cell.text:
                        itens_text = ""
                        for i, item in enumerate(itens, 1):
                            itens_text += f"{i}. {item}\n"
                        # Substituição na célula preservando parágrafos internos
                        for p in cell.paragraphs:
                            if '<Itens da OS>' in p.text:
                                p.text = p.text.replace('<Itens da OS>', itens_text.strip())
                                for run in p.runs:
                                    run.font.name = 'Inter'
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Nome do arquivo de saída
        filename = f"Proposta OS {ano_os}-{num_os} - {titulo_os}.docx"
        output_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
        doc.save(output_path)

        return output_path

    def generate_relatorio_entrega(self, autor, data_servidor, proposta_file, hu_files):
        self._ensure_runtime_folders()
        template_path = self._resolve_template_path('TEMPLATE_RELATORIO_ENTREGA_DOCX')

        proposta_dir, hu_dir, relatorio_dir = self._new_execution_paths()
        proposta_filename = self._sanitize_upload_name(proposta_file.filename, 'proposta.docx')
        proposta_target = proposta_dir / proposta_filename
        proposta_file.save(str(proposta_target))

        identity = self._extract_proposta_identity(proposta_target, proposta_file.filename)
        workbook = self._resolve_data_source_file()
        sheet_row = self._read_xlsx_row_by_os(workbook, identity['ano_os'], identity['os'])

        saved_hus = []
        for hu_file in hu_files:
            if not hu_file or not hu_file.filename:
                continue
            original_stem = Path(hu_file.filename).stem
            hu_name = self._sanitize_upload_name(hu_file.filename, f'hu_{len(saved_hus)+1}.docx')
            hu_target = hu_dir / hu_name
            hu_file.save(str(hu_target))
            saved_hus.append({'path': hu_target, 'label': original_stem})

        if not saved_hus:
            raise ValueError('Nenhuma HU válida foi enviada para gerar o relatório.')

        introducao = self._extract_intro_from_proposta_docx(proposta_target)
        data_atual = data_servidor or datetime.now().strftime('%d/%m/%Y')
        assunto = identity['assunto'] or sheet_row.get('assunto', '')
        inicio = self._excel_serial_to_date(sheet_row.get('inicio', ''))
        fim = self._excel_serial_to_date(sheet_row.get('fim', ''))
        hst = sheet_row.get('hst', '')
        valor = self._format_brl_value(float(sheet_row.get('valor', '0') or 0))

        doc = Document(template_path)
        hu_lines = [item['label'] for item in saved_hus]
        replacements = {
            '<Autor>': autor,
            '{{AUTOR}}': autor,
            '{{Autor}}': autor,
            '<Data atual do servidor>': data_atual,
            '{{DATA_ATUAL_SERVIDOR}}': data_atual,
            '{{Data do Servidor}}': data_atual,
            '<Introdução da Proposta>': introducao,
            '<Introducao da Proposta>': introducao,
            '{{INTRODUCAO_PROPOSTA}}': introducao,
            '{{Replicar Introdução da Proposta da Ordem de Serviço}}': introducao,
            '{{Replicar Introducao da Proposta da Ordem de Servico}}': introducao,
            '{{ANO_OS}}': identity['ano_os'],
            '{{Ano_OS}}': identity['ano_os'],
            '{{OS}}': identity['os'],
            '{{Assunto}}': assunto,
            '{{Início}}': inicio,
            '{{Inicio}}': inicio,
            '{{Fim}}': fim,
            '{{HST}}': str(hst),
            '{{Valor}}': valor,
        }
        self._write_hu_list_item_41(doc, saved_hus)
        self._replace_placeholders_everywhere(doc, replacements)
        self._normalize_item9_anexos_heading_position(doc)
        self._ensure_item9_anexos_starts_new_page(doc)
        self._ensure_two_blank_lines_after_item9_heading(doc)
        self._write_hu_item_9(doc, saved_hus)

        report_title = self._sanitize_output_filename(
            f"Relatório de Entrega OS {identity['ano_os']}-{identity['os']} - {assunto}".strip()
        )
        output_docx = relatorio_dir / f'{report_title}.docx'
        doc.save(str(output_docx))

        final_output = Path(current_app.config['UPLOAD_FOLDER']) / output_docx.name
        shutil.copy2(output_docx, final_output)
        return str(final_output), str(relatorio_dir), str(hu_dir)
